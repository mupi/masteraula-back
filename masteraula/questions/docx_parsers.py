from html.parser import HTMLParser

from docx import *
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib
from urllib.error import HTTPError

import re
import os
import datetime

import time

current_milli_time = lambda: int(round(time.time() * 1000))

class Question_Parser(HTMLParser):
    ''' Classe responsavel por realizer o parser para geracao de listas de
    questoes  '''
    # Controle do documento
    docx_title = ''
    questionCounter = 0
    question_item = 'a'
    year_source = ''
    page_width = 0

    # Controle do paragrafo
    paragraph = None
    run = None

    # Estilos
    underline = False
    italic = False
    bold = False
    subscript = False
    superscript = False
    font_size = 10

    # Flags
    alternatives = False         # resposta em uma unica linha
    is_table = False        # texto dentro da tabela
    is_list = False         # Listas de itens
    is_poem = False
    is_div_poem = False
    line_columns = True
    printed_counter = False

    # Constantes
    max_image_height = 0

    # Variaveis de controle de tabelas
    table = None
    cell_paragraph = None
    cell_run = None
    row_index = -1
    column_index = -1

    def __init__(self, title):
        HTMLParser.__init__(self)
        # Faz o novo documento
        self.document = Document()
        self.docx_title = title

        section = self.document.sections[0]

        # Tamanho
        section.page_height = Pt(842)
        section.page_width = Pt(595)

        section.left_margin = section.right_margin = Pt(50)
        section.top_margin = section.bottom_margin = Pt(50)

        self.page_width = section.page_width - (section.left_margin + section.right_margin)
        self.max_image_height = (int)((section.page_height-(section.top_margin + section.bottom_margin))/3)

    def parse_heading(self, list_header):
        '''Cabecalho do gabarito gerado, colocando o nome da lista e tambem
        a data que foi criada'''
        
        pass
    
        # doc = list_header
        
        # header = self.document.add_heading(doc.name, 1)
        # header.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        # header.paragraph_format.space_after = Pt(8)

        # table = self.document.add_table(rows=0, cols=1)
        # table.style = 'TableGrid'

        # if doc.institution_name:
        #     table.add_row().cells[0].text = 'Instituição: ' + doc.institution_name
        # if doc.discipline_name:
        #     table.add_row().cells[0].text = 'Disciplina: ' + doc.discipline_name
        # if doc.professor_name:
        #     table.add_row().cells[0].text = 'Professor(a): ' + doc.professor_name

        
        # if doc.student_indicator:
        #     table.add_row().cells[0].text ='Aluno:'
        # if doc.class_indicator:
        #     table.add_row().cells[0].text ='Turma:'
        # if doc.score_indicator:
        #     table.add_row().cells[0].text ='Nota da avaliação:'
        # if doc.date_indicator:
        #     table.add_row().cells[0].text ='Data:      /      /      '
      
    def parse_list_questions(self, list_questions, resolution):
        '''Faz o parser de cada questao da lista, tambem chamara a funcao
        auxiliar que faz o parser das respostas de cada questao'''
        for question in list_questions:
            # Titulo de cada questao
            self.reset_flags()
            self.start_question()

            self.year_source = str(self.questionCounter)
            if (question.source != None and question.source != "") and question.year != None:
                self.year_source = self.year_source + ' (' + str(question.year)  + ' - '+ question.source + ') '
            elif (question.source != None and question.source != ""):
                self.year_source = self.year_source + ' (' + str(question.source) + ') '
            elif question.year != None:
                self.year_source = self.year_source + ' (' + str(question.year) + ') '

            self.paragraph = None
            self.run = None

            # o WysWyg adiciona varios \r, o parser nao trata esse caso especial entao remove-se todas suas ocorrencias
            self.feed(question.statement.replace('\n', ' ').replace('\r', ''))

            # Respotas enumeradas de a a quantidade de respostas
            self.parse_list_alternatives(question.alternatives.all())

            if resolution:
                self.parse_resolution(question)

    def end_parser(self):
        '''Finaliza o parser excluindo todos os paragrafos que eventualmente
        ficaram em branco'''
        # for para in self.document.paragraphs:
        #     if para.text == "":
        #         self.delete_paragraph(para)

        self.document.add_page_break()
        self.document.save(self.docx_title)

# Auxiliar functions

    def parse_resolution(self, question):
        '''Faz o parser das resolucoes de cada questao'''
        resolution = question.resolution
        question_item = 'a'
        answered = False
        # Procura pela questao correta
        for answer in question.alternatives.all():
            if answer.is_correct:
                answered = True
            elif not answered:
                question_item = chr(ord(question_item) + 1)

        if answered:
            self.init_parser_answer()
            self.bold = True
            self.feed('Resposta: ' + question_item)

        if resolution is not None and resolution != '':
            self.init_parser_answer()
            self.bold = True
            self.feed(resolution.replace('\n', '').replace('\r', ''))


    def parse_list_alternatives(self, list_answer):
        '''Faz o parser das respostas de cada questao'''
        # Seta a flag para escrever as resposas
        self.alternatives = True
        for answer in list_answer:
            self.init_parser_answer()

            to_parse = self.question_item + ') ' + answer.text
            self.feed(to_parse.replace('\t', '').replace('\r', '').replace('\n',''))
            self.question_item = chr(ord(self.question_item) + 1)

    def reset_flags(self):
        self.underline = False
        self.italic = False
        self.bold = False
        self.subscript = False
        self.superscript = False

    def init_parser_answer(self):
        '''Prepara o parser zerando todas as variaveis de estilo e justificando
        o texto'''
        self.paragraph = self.document.add_paragraph()
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.paragraph.paragraph_format.space_after = Pt(0)
        self.paragraph.paragraph_format.space_before = Pt(0)
        self.run = self.paragraph.add_run()
        self.font_size = 10
        self.run.font.size = Pt(self.font_size)

        self.reset_flags()

    def delete_paragraph(self, paragraph):
        '''Deleta o paragrafo passado como parametro'''
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def start_question(self):
        '''Prepara a questao para ser tratada, zerando algumas variaveis de
        controle, colocando o cabecalho e aumentando o numero de questoes'''
        self.questionCounter = self.questionCounter + 1

        self.document.add_heading(level=2)

        # Flag de respostas False pois comecara com o enunciado
        self.alternatives = False
        self.printed_counter = False
        self.question_item = 'a'

# Parser Methods

    def handle_starttag(self, tag, attrs):
        # Adiciona imagem
        if tag == 'img':
            if self.run == None:
                self.run = self.add_or_create_run()

            if not self.printed_counter:
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self.run.font.size = Pt(10)
                self.run.add_text(self.year_source)
                self.printed_counter = True
                self.paragraph = self.document.add_paragraph()
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.run = self.add_or_create_run()

            # Variaveis das imagens
            height = None
            width = None
            src = None
            # sai a procura das variaveis das imagens
            for attr in attrs:
                if attr[0] == 'src':
                    src = attr[1]
            # Faz o download da imagem
            try:
                self.paragraph = self.document.add_paragraph()
                self.run = self.paragraph.add_run()
                image_name = self.docx_title + str(current_milli_time()) + ".png"
                urllib.request.urlretrieve(src, image_name)
                image = self.run.add_picture(image_name)
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                os.remove(image_name)

                if  not self.alternatives:
                    self.paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if (image.width > self.page_width):
                    widthMult = self.page_width / image.width
                    image.width = self.page_width
                    image.height = (int)(image.height * widthMult)

                if (image.height > self.max_image_height):
                    heightMult = self.max_image_height / image.height
                    image.height = self.max_image_height
                    image.width = (int)(image.width * heightMult)

            except HTTPError as e:
                if e.code == 403:
                    p = self.document.add_paragraph()
                    if  not self.alternatives:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run('Imagem não encontrada!')
                    r.font.size = Pt(8)
                    r.bold = True
            
        # Habilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        elif tag == 'u' or tag == 'em' or tag == 'strong' or tag == 'sub' or tag == 'sup':
            self.run = self.add_or_create_run()
            self.run.font.size = Pt(self.font_size)
            if tag == 'u':
                self.underline = True
            elif tag == 'em':
                self.italic = True
            elif tag == 'strong':
                self.bold = True
            elif tag == 'sub':
                self.subscript = True
            elif tag == 'sup':
                self.superscript = True

        # listas
        elif tag =='li':
            self.paragraph = self.document.add_paragraph(style='ListBullet')
            self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.run = self.paragraph.add_run()
            self.font_size = 10
            self.run.font.size = Pt(self.font_size)

        elif tag == 'h1':
            self.paragraph = self.document.add_paragraph()
            self.paragraph.paragraph_format.space_after = Pt(0)
            self.paragraph.paragraph_format.space_before = Pt(0)
            self.run = self.paragraph.add_run()
            self.font_size = 12
            self.run.font.size = Pt(self.font_size)

        # Pula linha
        if tag == 'br':
            if self.run == None:
                self.run = self.add_or_create_run()

            if self.is_poem or self.is_div_poem:
                self.run.add_text("\n")
            else:
                self.paragraph = self.document.add_paragraph()
                self.paragraph.paragraph_format.space_after = Pt(0)
                self.paragraph.paragraph_format.space_before = Pt(0)
                self.run = self.paragraph.add_run()
                self.run.font.size = Pt(self.font_size)
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Textos de tabelas
        elif tag == 'p' and self.is_table:
            self.cell_paragraph = self.table.cell(
                        self.row_index, self.column_index).add_paragraph()
            self.cell_run = self.cell_paragraph.add_run()

        # Textos normais, nota que textos de respostas ocupam somente uma linha,
        # mesmo que contenham varias tags <p>
        elif tag == 'p' and not self.alternatives:
            self.paragraph = self.document.add_paragraph()
            self.paragraph.paragraph_format.space_after = Pt(0)
            self.paragraph.paragraph_format.space_before = Pt(0)
            self.run = self.paragraph.add_run()
            self.font_size = 10
            self.run.font.size = Pt(self.font_size)

            # Verificacao da classe
            p_class = ""
            for attr in attrs:
                if attr[0] == 'class':
                    p_class = attr[1]
            if p_class == "question_source":
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self.font_size = 8
                self.run.font.size = Pt(self.font_size)
            elif p_class == "estrofe" or p_class == "poema":
                self.is_poem = True
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif tag == 'div' and not self.alternatives:
            # Verificacao da classe
            p_class = ""
            for attr in attrs:
                if attr[0] == 'class':
                    p_class = attr[1]
            if p_class == "poema":
                self.is_div_poem = True

        # Controle de tabelas
        elif tag == 'table':
            self.is_table = True
            self.table = self.document.add_table(rows=0, cols=0)
            self.table.style = 'TableGrid'
        elif tag == 'tr':
            self.table.add_row()
            self.row_index = self.row_index + 1
            self.column_index = -1
        elif tag == 'td':
            # Controle de adicao de tabelas, atualmente nao possui table merge
            if self.line_columns:
                self.table.add_column(1000)     # possivel bug do python-docx
            self.column_index = self.column_index + 1

    def handle_endtag(self, tag):
        # Desabilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        if tag == 'u' or tag == 'em' or tag == 'strong' or tag == 'sub' or tag == 'sup':
            self.run = self.paragraph.add_run()
            self.run.font.size = Pt(self.font_size)
            if tag == 'u':
                self.underline = False
            elif tag == 'em':
                self.italic = False
            elif tag == 'strong':
                self.bold = False
            elif tag == 'sub':
                self.subscript = False
            elif tag == 'sup':
                self.superscript = False


        elif tag == 'h1':
            self.paragraph = None
            self.run = None

        # Desabilita os textos
        elif tag == 'p' and self.is_table:
            self.cell_run = None
            self.cell_paragraph = None

        elif tag == 'p':
            self.run = None
            self.paragraph = None
            self.is_poem = self.is_div_poem

        elif tag == 'div':
            self.run = None
            self.paragraph = None
            self.is_div_poem = False

        elif tag == 'li':
            self.run = None
            self.paragraph = None

        # Desabilita as tabelas
        elif tag == 'table':
            self.is_table = False
        elif tag == 'tr':
            self.line_columns = False
        elif tag == 'td':
            self.cell_run = None
            # Exclui paragrafos vazios na celula atual, workaround encontrado
            # para melhorar o visual da lista gerada
            cell = self.table.cell(self.row_index, self.column_index)
            for para in cell.paragraphs:
                if para.text == "":
                    self.delete_paragraph(para)

    def handle_data(self, data):
        # Escreve o texto na tabela, verifica se o dado contem alguma coisa
        # importante, (existem linhas com tabulacoes que nao devem ser postas)
        if self.is_table and re.sub(r"\W", "", data) != "":
            # Texto normal sem tag
            if self.cell_run is None:
                self.cell_paragraph = self.table.cell(self.row_index,
                                self.column_index).add_paragraph()
                self.cell_paragraph.add_run(data)
            # Texto com tags
            else:
                self.cell_run.add_text(data)

        # Escreve o texto normalmente
        elif self.run is not None:
            if not self.printed_counter:
                self.run.font.size = Pt(10)
                self.run.add_text(self.year_source + ' ')

                self.printed_counter = True

                self.run = self.paragraph.add_run()
                self.run.font.size = Pt(self.font_size)

            self.run.add_text(data)

            font = self.run.font
            font.underline = self.underline
            font.italic = self.italic
            font.bold = self.bold
            font.subscript = self.subscript
            font.superscript = self.superscript

        # Escreve o texto normalmente
        elif data.replace(" ","") != "":
            self.run = self.add_or_create_run()
            self.font_size = 10
            self.run.font.size = Pt(self.font_size)

            if self.is_poem or self.is_div_poem:
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if self.printed_counter:
                self.run.add_text(data)
            else:
                self.run.add_text(self.year_source + " " + data)
                self.printed_counter = True

    def parse_alternatives_list_questions(self, list_questions):
        '''Faz o parser do gabarito, em formato de tabela'''

        question_counter = 0

        # Adiciona a tabela e seus cabecalhos
        self.document.add_page_break()
        self.document.add_heading("Gabarito")

        table = self.document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Questão'
        hdr_cells[1].text = 'Resposta'

        # Faz o parser de cada questao
        for question in list_questions:
            question_counter = question_counter + 1
            table.add_row()
            table.cell(question_counter, 0).text = str(question_counter)

            question_item = 'a'
            answered = False
            # Procura pela questao correta
            for answer in question.alternatives.all():
                if answer.is_correct:
                    table.cell(question_counter, 1).text = question_item
                    answered = True
                question_item = chr(ord(question_item) + 1)

            # Se nao houve questao correta, fala sem
            # resposta
            if not answered:
                table.cell(question_counter, 1).text = 'Sem resposta'


    def add_or_create_run(self):
        if self.paragraph == None:
            self.paragraph = self.document.add_paragraph()
            self.paragraph.paragraph_format.space_after = Pt(0)
            self.paragraph.paragraph_format.space_before = Pt(0)
        return self.paragraph.add_run()

class Answer_Parser():
    # Controle do documento
    document = None
    docx_title = None

    def __init__(self, title):
        self.document = Document()
        self.docx_title = title

    def parse_alternatives_list_questions(self, list_questions):
        '''Faz o parser do gabarito, em formato de tabela'''

        question_counter = 0

        # Adiciona a tabela e seus cabecalhos
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Questao'
        hdr_cells[1].text = 'Resposta'

        # Faz o parser de cada questao
        for question in list_questions:
            question_counter = question_counter + 1
            table.add_row()
            table.cell(question_counter, 0).text = str(question_counter)

            question_item = 'a'
            answered = False
            # Procura pela questao correta
            for answer in question.alternatives.all():
                if answer.is_correct:
                    table.cell(question_counter, 1).text = question_item
                    answered = True
                question_item = chr(ord(question_item) + 1)

            # Se nao houve questao correta, fala sem
            # resposta
            if not answered:
                table.cell(question_counter, 1).text = 'Sem resposta'
