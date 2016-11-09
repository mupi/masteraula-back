# -*- coding: utf-8 -*-

from django.views.generic import DetailView, ListView, RedirectView, UpdateView, TemplateView, CreateView, DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect, Http404
from django.db.models import Q
from django.shortcuts import redirect
from django.urls import reverse

from html.parser import HTMLParser

from docx import *
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import json
import datetime
import io
import re

from .models import Question, Question_List, QuestionQuestion_List

class QuestionDetailView(LoginRequiredMixin, DetailView):
    model = Question
    template_name = "questions/question_detail.html"
    context_object_name = "question"

class QuestionListView(LoginRequiredMixin, ListView):
    model = Question
    queryset = Question.objects.order_by('pk')
    template_name = "questions/question_list.html"
    context_object_name = "question_list"
    paginate_by = 10



class Question_ListDetailView(LoginRequiredMixin, DetailView):
    model = Question_List
    template_name = "questions/question_list_detail.html"
    context_object_name = "question_list"

    # Verifica se a lista eh privada e se for, somente mostrar para seu dono
    def dispatch(self, request, *args, **kwargs):
        question_list = self.get_object()
        if question_list.private and question_list.owner != self.request.user:
            return redirect(reverse('questions:question_list_list'))
        return super(Question_ListDetailView, self).dispatch(request, *args, **kwargs)

class Question_ListDeleteView(LoginRequiredMixin, DeleteView):
    model = Question_List
    context_object_name = "question_list"
    success_url = "/questions/question_lists"

    # Verifica se o request.user eh o dono da lista a ser deletada
    def dispatch(self, request, *args, **kwargs):
        question_list = self.get_object()
        if question_list.owner != self.request.user:
            return redirect(reverse('questions:question_list_detail', args=(self.kwargs['pk'],)))
        return super(Question_ListDeleteView, self).dispatch(request, *args, **kwargs)

class Question_ListCreateView(LoginRequiredMixin, CreateView):
    model = Question_List
    fields = ['question_list_header', 'private']
    template_name = "questions/question_list_create.html"

    # Necessario para passar informacoes para o front-end disponibilizar quais
    # exercicios foram selecionados, mosrtando no check-box
    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListCreateView, self).get_context_data(**kwargs)
        # Atualiza com base no valor salvo na sessao (ou um novo, caso nao tenha)
        if not 'checked_questions' in self.request.session or not self.request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = self.request.session['checked_questions']

        # Torna-o acessivel diretamente no template e feito dessa forma para
        # garantir a ordem de selecao
        context['checked_questions'] = [Question.objects.get(pk=question_id)
                                        for question_id in checked_questions ]

        return context

    def form_valid(self, form):
        new_list = form.save(commit=False)
        new_list.owner = self.request.user
        # Salva a primeira vez para se ter um PK e assim poder fazer a relacao
        # many-to-many das questoes da lista
        new_list.save()

        if not 'checked_questions' in self.request.session or not self.request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = self.request.session['checked_questions']

        # Faz as relacoes many-to-many respeitando a ordem
        questions = [Question.objects.get(pk=question_id)
                    for question_id in checked_questions ]
        order = 1

        for question in questions:
            relation = QuestionQuestion_List(question=question,
                                            question_list=new_list, order=order)
            order = order + 1
            relation.save()

        # Apaga o valor da session para evitar a dados na criacao de outra lista
        self.request.session['checked_questions'] = []
        return HttpResponseRedirect(reverse('questions:question_list_detail', args=(str(new_list.pk),)))

class Question_ListEditView(LoginRequiredMixin, UpdateView):
    model = Question_List
    fields = ['question_list_header', 'private']
    template_name = "questions/question_list_edit.html"

    # Verifica se o request.user eh o dono da lista a ser editada
    def dispatch(self, request, *args, **kwargs):
        question_list = self.get_object()
        if question_list.owner != self.request.user:
            return redirect(reverse('questions:question_list_detail', args=(self.kwargs['pk'],)))
        return super(Question_ListEditView, self).dispatch(request, *args, **kwargs)

    # Somente adiciona as novas questoes, para remover eh o Question_ListRemoveQuestionsView
    def form_valid(self, form):
        new_list = form.save(commit=False)

        if 'checked_edit_add_questions' in self.request.session and self.request.session['checked_edit_add_questions']:
            checked_questions = self.request.session['checked_edit_add_questions']
            questions = [Question.objects.get(pk=question_id)
                            for question_id in checked_questions ]

            # Relacoes many-to-many das novas questoes a serem adicionadas
            order = len(new_list.questions.all()) + 1
            for question in questions:
                relation = QuestionQuestion_List(question=question,
                                                question_list=new_list, order=order)
                order = order + 1
                relation.save()
            new_list.save()

            # Apaga o valor para mostrar dados consistentes no front-end e evitar
            # a propagacao dos dados dessa lista na edicao de outras
            self.request.session['checked_edit_add_questions'] = []

        return HttpResponseRedirect(reverse('questions:question_list_detail', args=(str(new_list.pk),)))

    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListEditView, self).get_context_data(**kwargs)

        # Verifica para mostrar somente dados referentes a atualizacao
        # da mesma lista e os dados que estao no request
        if not 'checked_edit_add_questions' in self.request.session or not self.request.session['checked_edit_add_questions']:
            # Zera todas as variaveis de controle para a insercao de novas questoes
            self.request.session['question_list_edit_id'] = None        # ID da lista
            self.request.session['checked_edit_add_questions'] = []     # Questoes a serem adicionadas
            self.request.session['checked_edit_questions'] = []         # Questoes ja adicionadas selecionadas (para exclusao)
        elif self.request.session['question_list_edit_id'] == self.kwargs['pk']:
            checked_questions = [Question.objects.get(pk=question_id) for question_id in self.request.session['checked_edit_add_questions']]
            context['checked_edit_add_questions'] = checked_questions

        return context

class Question_ListRemoveQuestionsView(LoginRequiredMixin, UpdateView):
    model = Question_List
    fields = []

    # Verifica se o request.user eh o dono da lista a ser editada
    def dispatch(self, request, *args, **kwargs):
        question_list = self.get_object()
        if question_list.owner != self.request.user:
            return redirect(reverse('questions:question_list_detail', args=(self.kwargs['pk'],)))
        return super(Question_ListRemoveQuestionsView, self).dispatch(request, *args, **kwargs)

    # Somente remove as questoes, para adicionar eh o Question_ListEditView
    def form_valid(self, form):
        edited_list = form.save(commit=False)

        # Remove questao a questao das questoes ja adicionadas selecionadas
        if 'checked_edit_questions' in self.request.session and self.request.session['checked_edit_questions']:
            checked_questions = self.request.session['checked_edit_questions']
            questions = edited_list.questions.filter(pk__in=checked_questions)

            for question in questions:
                QuestionQuestion_List.objects.get(question=question, question_list = edited_list).delete()
            edited_list.save()

            # Apaga o valor para mostrar dados consistentes no front-end e evitar
            # a propagacao dos dados dessa lista na edicao de outras
            self.request.session['checked_edit_questions'] = []

        return HttpResponseRedirect(reverse('questions:question_list_edit', args=(str(edited_list.pk),)))

class Question_ListEditListView(LoginRequiredMixin, ListView):
    model = Question
    template_name = "questions/question_list_edit_list.html"
    context_object_name = "question_list"
    paginate_by = 10

    # Verifica se o request.user eh o dono da lista a ser editada
    def dispatch(self, request, *args, **kwargs):
        question_list = Question_List.objects.get(id=self.kwargs['pk'])
        if question_list.owner != self.request.user:
            return redirect(reverse('questions:question_list_detail', args=(self.kwargs['pk'],)))
        return super(Question_ListEditListView, self).dispatch(request, *args, **kwargs)

    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListEditListView, self).get_context_data(**kwargs)

        # Verifica se a session esta armazenando dados de edicao da lista de
        # questoes que esta sendo atualizada no momento, evitando dados mortos
        # na session
        if not self.request.session['question_list_edit_id'] or self.request.session['question_list_edit_id'] != self.kwargs['pk']:
            self.request.session['checked_edit_add_questions'] = []
        self.request.session['question_list_edit_id'] = self.kwargs['pk']

        list_exclude_questions = Question_List.objects.get(id=self.kwargs['pk']).questions.all()

        context['question_list'] = Question.objects.exclude(pk__in=list_exclude_questions)
        # Armazena para fazer a verificaçao acima
        context['pk_slug'] = self.kwargs['pk']

        return context

class Question_ListListView(LoginRequiredMixin, ListView):
    model = Question_List
    template_name = "questions/question_list_list.html"
    context_object_name = "question_list_list"
    success_url = "/questions"

    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListListView, self).get_context_data(**kwargs)

        context['question_list_list'] = Question_List.objects.filter(Q(private=False) | Q(owner=self.request.user.id))
        return context

class Question_ListOwnListView(LoginRequiredMixin, ListView):
    model = Question_List
    template_name = "questions/question_list_own_list.html"
    context_object_name = "question_list_list"
    success_url = "/questions"

    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListOwnListView, self).get_context_data(**kwargs)

        context['question_list_list'] = Question_List.objects.filter(owner=self.request.user.id)
        return context

class Question_ListCloneView(LoginRequiredMixin, CreateView):
    model = Question_List
    fields = ['question_list_header', 'private']
    template_name = "questions/question_list_clone.html"
    context_object_name = "question_list"

    def get_context_data(self, *args, **kwargs):
        context = super(Question_ListCloneView, self).get_context_data(**kwargs)

        context['cloned_from_list'] = Question_List.objects.get(id=self.kwargs['pk'])
        return context

    # Clona a lista, copiando as questoes mas alterando o dono e a origem da copia
    def form_valid(self, form):
        # Dados da lista original
        cloned_from_list = self.get_context_data()['cloned_from_list']
        # Dados da nova lista (nome e privado)
        new_list = form.save(commit=False)
        new_list.owner = self.request.user
        new_list.cloned_from = cloned_from_list
        # Faz o primeiro save para se poder fazer as relacoes many-to-many
        new_list.save()

        #Adiciona a relacao many-to-many uma a uma
        if not cloned_from_list or not cloned_from_list.questions.all:
            questions = []
        else:
            questions = cloned_from_list.questions.all()

        for question in questions:
            new_list.questions.add(question)
        new_list.save()

        self.request.session['cloned_from_list'] = None
        return HttpResponseRedirect(reverse('questions:question_list_detail', args=(str(new_list.pk),)))


# Trata quando um usuario seleciona o checkbox para fazer uma nova lista nas
# listas de questoes, armazenando tal informacao na sessao
def check_question(request):
    if (request.method == 'POST'):
        questionPk = request.POST.get('questionPK')
        checked = request.POST.get('checked')

        # Pega a lista ja salva na sessao ou cria uma caso nao tenha sido ainda
        if not 'checked_questions' in request.session or not request.session['checked_questions']:
            checked_questions = []
        else:
            checked_questions = request.session['checked_questions']

        # Adiciona se foi selecionada o checkbox ou remove caso tenha sido deselecionado
        if checked == 'true':
            checked_questions.append(int(questionPk))
        else:
            checked_questions = [item for item in checked_questions if item != int(questionPk)]

        request.session['checked_questions'] = checked_questions

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    raise Http404("Method GET not allowed in check_question!")

# Limpa todas as questoes que foram selecionadas usando o metodo acima
def clear_questions(request):
    if (request.method == 'POST'):

        request.session['checked_questions'] = []

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    raise Http404("Method GET not allowed in check_question!")

# Trata quando um usuario seleciona o checkbox para editar questoes ja pertencentes as
# listas de questoes, armazenando tal informacao na sessao
def check_question_edit_list(request):
    if (request.method == 'POST'):
        questionPk = request.POST.get('questionPK')
        checked = request.POST.get('checked')

        # Pega a lista ja salva na sessao ou cria uma caso nao tenha sido ainda
        if not 'checked_edit_questions' in request.session or not request.session['checked_edit_questions']:
            checked_questions = []
        else:
            checked_questions = request.session['checked_edit_questions']

        # Adiciona se foi selecionada o checkbox ou remove caso tenha sido deselecionado
        if checked == 'true':
            checked_questions.append(int(questionPk))
        else:
            checked_questions = [item for item in checked_questions if item != int(questionPk)]

        request.session['checked_edit_questions'] = checked_questions

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    raise Http404("Method GET not allowed in check_question!")

# Trata quando um usuario seleciona o checkbox para adicionar questoes as
# listas de questoes, armazenando tal informacao na sessao
def check_question_edit_add_list(request):
    if (request.method == 'POST'):
        questionPk = request.POST.get('questionPK')
        checked = request.POST.get('checked')

        # Pega a lista ja salva na sessao ou cria uma caso nao tenha sido ainda
        if not 'checked_edit_add_questions' in request.session or not request.session['checked_edit_add_questions']:
            checked_questions = []
        else:
            checked_questions = request.session['checked_edit_add_questions']

        # Adiciona se foi selecionada o checkbox ou remove caso tenha sido deselecionado
        if checked == 'true':
            checked_questions.append(int(questionPk))
        else:
            checked_questions = [item for item in checked_questions if item != int(questionPk)]

        request.session['checked_edit_add_questions'] = checked_questions

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    raise Http404("Method GET not allowed in check_question!")

# Limpa todas as questoes que foram selecionadas usando o metodo acima
def clear_questions_edit_add_list(request):
    if (request.method == 'POST'):

        request.session['checked_edit_add_questions'] = []

        return HttpResponse(
            json.dumps({"status" : "success"}),
            content_type="application/json"
        )
    raise Http404("Method GET not allowed in check_question!")

class QuestionCreate(CreateView):
    model = Question
    fields = [
        'question_header',
        'question_text',
        'resolution',
        'level',
        'author',
        'tags',
    ]

class Question_HeaderParser(HTMLParser):
    ''' Classe responsavel por realizer o parser para geracao de listas de
    questoes  '''
    # Estilo do texto (sublinhado, negrito e italico)
    underline = False
    bold = False
    italic = False

    # Controle do paragrafo
    paragraph = None
    run = None

    # Flags
    respostas = False   # resposta em uma unica linha
    is_table   = False     # texto dentro da tabela
    line_columns = True

    # Variaveis de controle de tabelas
    table = None
    cell_paragraph = None
    cell_run = None
    row_index = -1
    column_index = -1

    def __init__(self, document):
        HTMLParser.__init__(self)
        self.document = document

    def handle_starttag(self, tag, attrs):
        # Adiciona imagem
        if tag == 'img':
            # Variaveis das imagens
            height = None
            width = None
            src = None
            # sai a procura das variaveis das imagens
            for attr in attrs:
                if attr[0] == 'src':
                    src = 'mupi_question_database' + attr[1]
                elif attr[0] == 'style':
                    values = dict(item.split(":") for item in attr[1].split(";"))
                    for value in values:
                        if value.replace(' ', '') == 'width':
                            width = re.sub('\D', '', values[value])
                        elif value.replace(' ', '') == 'height':
                            height = re.sub('\D', '', values[value])
            # Coloca a imagem de acordo com sua respectiva dimensao definida (supondo DPI = 180)
            if width:
                self.run.add_picture(src, width=Inches(int(width)/180))
            else:
                self.run.add_run().add_picture(src)

        # Habilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        elif tag == 'u':
            self.underline = True
        elif tag == 'em':
            self.italic = True
        elif tag == 'strong':
            self.bold = True

        # Textos de tabelas
        elif tag == 'p' and self.is_table:
            self.cell_paragraph = self.table.cell(
                        self.row_index, self.column_index).add_paragraph()
            self.cell_run = self.cell_paragraph.add_run()

        # Textos normais, nota que textos de respostas ocupam somente uma linha,
        # mesmo que contenham varias tags <p>
        elif tag == 'p' and not self.respostas:
            self.paragraph = self.document.add_paragraph()
            self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.run = self.paragraph.add_run()

        # Controle de tabelas
        elif tag == 'table':
            self.is_table = True
            self.table = self.document.add_table(rows=0, cols=0)
            self.table.style = 'TableGrid'
        elif tag == 'tr':
            self.table.add_row()
            self.row_index = self.row_index + 1
            self.column_index = -1
        elif tag == 'td':
            # Controle de adicao de tabelas, atualmente nao possui table merge
            if self.line_columns:
                self.table.add_column(1000)     # possivel bug do python-docx
            self.column_index = self.column_index + 1

    def handle_endtag(self, tag):
        # Desabilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        if tag == 'u':
            self.underline = False
        elif tag == 'em':
            self.italic = False
        elif tag == 'strong':
            self.bold = False

        # Desabilita os textos
        elif tag == 'p' and self.is_table:
            self.cell_run = None

        elif tag == 'p':
            self.run = None

        # Desabilita as tabelas
        elif tag == 'table':
            self.is_table = False
        elif tag == 'tr':
            self.line_columns = False
        # Desabilita as tabelas
        elif tag == 'td':
            self.cell_run = None
            # Exclui paragrafos vazios na celula atual, workaround encontrado
            # para melhorar o visual da lista gerada
            cell = self.table.cell(self.row_index, self.column_index)
            for para in cell.paragraphs:
                if para.text == "":
                    self.delete_paragraph(para)

    def handle_data(self, data):
        # Escreve o texto na tabela, verifica se o dado contem alguma coisa
        # importante, (existem linhas com tabulacoes que nao devem ser postas)
        if self.is_table and re.sub(r"\W", "", data) != "":
            # Texto normal sem tag
            if self.cell_run is None:
                self.cell_paragraph = self.table.cell(self.row_index,
                                self.column_index).add_paragraph()
                self.cell_paragraph.add_run(data)
            # Texto com tags
            else:
                self.cell_run.add_text(data)
                font = self.cell_run.font
                font.italic     = self.italic
                font.underline  = self.underline
                font.bold       = self.bold

        # Escreve o texto normalmente
        elif self.run is not None:
            self.run.add_text(data)
            font = self.run.font
            font.italic     = self.italic
            font.underline  = self.underline
            font.bold       = self.bold

    def init_parser(self, paragraph):
        # Reseta todas as variaveis responsaveis pela escrita do enunciado
        self.paragraph = paragraph
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.run = self.paragraph.add_run()
        self.underline = False
        self.bold = False
        self.italic = False

    def init_parser_new(self):
        # Reseta todas as variaveis responsaveis pela escrita do enunciado
        self.paragraph = self.document.add_paragraph()
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.run = self.paragraph.add_run()
        self.underline = False
        self.bold = False
        self.italic = False

    def end_parser(self):
        '''Finaliza o parser excluindo todos os paragrafos que eventualmente
        ficaram em branco'''
        for para in self.document.paragraphs:
            if para.text == "":
                self.delete_paragraph(para)

    def set_respostas(self):
        '''Seta a flag para tratamento de respostas, que devem contem uma linha
        apenas '''
        self.respostas = True

    def delete_paragraph(self, paragraph):
        '''Funcao auxiliar que deleta determinado paragrafo passado como
        parametro a funcao'''
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


# Gera um arquivo .docx contendo a lista de exercicios selecionadas
def list_generator(request):
    # Metodo GET gera a lista preparada pelo POST e realiza seu download
    if (request.method == 'GET'):
        if not 'generate_list_questions' in request.session or not request.session['generate_list_questions']:
            raise Http404("Cant generate a null list.")
        list_questions = [Question.objects.get(pk=question_id)
                            for question_id in request.session['generate_list_questions'] ]
        if not list_questions:
            raise Http404("Cant generate an empty list doccument.")

        questionCounter = 1

        # Faz o novo documento
        document = Document()
        parser = Question_HeaderParser(document)
        docx_title="Test_List.docx"

        # Cabecalho da lista gerada
        document.add_heading(request.session['generate_list_header'])

        document.add_paragraph(
            "Lista gerada em {:s} as {:s}.".format(
                datetime.date.today().strftime('%d/%m/%Y'),
                datetime.datetime.today().strftime('%X')
            )
        )

        for question in list_questions:
            # Titulo de cada questao
            document.add_heading('Questao ' + str(questionCounter), level=2)
            p = document.add_paragraph('(')
            p.add_run(question.question_header).bold = True
            p.add_run(') ')
            # Passa o paragrafo do titulo da questao para conter o enunciado do lado do Titulo
            # Exemplo (Enem 2015) Enunciado:
            parser.init_parser(p)

            # o WysWyg adiciona varios \r, o parser nao trata esse caso especial entao remove-se todas suas ocorrencias
            parser.feed(question.question_text.replace('\r\n\t', ''))

            # Respotas enumeradas de a a quantidade de respostas
            itemChar = 'a'
            parser.set_respostas()
            for answer in question.answers.all():
                parser.init_parser_new()
                to_parse = itemChar + ') ' + answer.answer_text
                parser.feed(to_parse.replace('\r\n\t', ''))
                itemChar = chr(ord(itemChar) + 1)

            questionCounter = questionCounter + 1

        document.add_page_break()
        parser.end_parser()

        document.save(docx_title)
        data = open(docx_title, "rb").read()

        # Zera variavel para evitar multiplos acessos a ele
        request.session['generate_list_questions'] = None
        request.session['generate_list_header'] = None

        response = HttpResponse(
            data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        return response

    # Metodo so responde para GET e POST
    if (request.method != 'POST'):
        return HttpResponse(
            json.dumps({"status" : "error"}),
            content_type="application/json"
        )

    # Prepara a lista para o metodo GET
    list_header = request.POST.getlist('listHeader')
    questionsId = request.POST.getlist('questionsId[]')

    if questionsId == None or len(questionsId) == 0:
        return HttpResponse(
            json.dumps({"status" : "error"}),
            content_type="application/json"
        )
    request.session['generate_list_questions'] = questionsId
    request.session['generate_list_header'] = list_header

    return HttpResponse(
        json.dumps({'status' : 'ready'}),
        content_type="application/json"
    )
