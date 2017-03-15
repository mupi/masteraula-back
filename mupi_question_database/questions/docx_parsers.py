from html.parser import HTMLParser

from docx import *
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib

import re
import os
import datetime

class Question_Parser(HTMLParser):
    ''' Classe responsavel por realizer o parser para geracao de listas de
    questoes  '''
    # Controle do documetno
    docx_title = ''
    questionCounter = 0
    question_item = 'a'

    # Estilo do texto (sublinhado, negrito e italico)
    underline = False
    bold = False
    italic = False

    # Controle do paragrafo
    paragraph = None
    run = None

    # Flags
    respostas = False   # resposta em uma unica linha
    is_table   = False     # texto dentro da tabela
    line_columns = True

    # Variaveis de controle de tabelas
    table = None
    cell_paragraph = None
    cell_run = None
    row_index = -1
    column_index = -1

    def __init__(self, title):
        HTMLParser.__init__(self)
        # Faz o novo documento
        self.document = Document()
        self.docx_title = title

    def parse_heading(self, list_header):
        '''Cabecalho do gabarito gerado, colocando o nome da lista e tambem
        a data que foi criada'''
        self.document.add_heading(list_header)

        self.document.add_paragraph(
            "Lista gerada em {:s} as {:s}.".format(
                datetime.date.today().strftime('%d/%m/%Y'),
                datetime.datetime.today().strftime('%X')
            )
        )

    def parse_list_questions(self, list_questions):
        '''Faz o parser de cada questao da lista, tambem chamara a funcao
        auxiliar que faz o parser das respostas de cada questao'''
        for question in list_questions:
            # Titulo de cada questao
            self.start_question()
            self.init_parser()

            # o WysWyg adiciona varios \r, o parser nao trata esse caso especial entao remove-se todas suas ocorrencias
            self.feed(question.question_statement.replace('\r\n\t', ''))

            # Respotas enumeradas de a a quantidade de respostas
            self.parse_list_answers(question.answers.all())

    def end_parser(self):
        '''Finaliza o parser excluindo todos os paragrafos que eventualmente
        ficaram em branco'''
        # for para in self.document.paragraphs:
        #     if para.text == "":
        #         self.delete_paragraph(para)

        self.document.add_page_break()
        self.document.save(self.docx_title)

# Auxiliar functions

    def parse_list_answers(self, list_answer):
        '''Faz o parser das respostas de cada questao'''
        # Seta a flag para escrever as resposas
        self.respostas = True
        for answer in list_answer:
            self.init_parser_new()
            to_parse = self.question_item + ') ' + answer.answer_text
            self.feed(to_parse.replace('\r\n\t', ''))
            self.question_item = chr(ord(self.question_item) + 1)

    def init_parser(self):
        '''Prepara o parser zerando todas as variaveis de estilo e justificando
        o texto'''
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.run = self.paragraph.add_run()
        self.underline = False
        self.bold = False
        self.italic = False

    def init_parser_new(self):
        '''Prepara o parser criando um novo paragrafo, zerando todas as
        variaveis de estilo e justificando o texto'''
        self.paragraph = self.document.add_paragraph()
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.run = self.paragraph.add_run()
        self.underline = False
        self.bold = False
        self.italic = False

    def delete_paragraph(self, paragraph):
        '''Deleta o paragrafo passado como parametro'''
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def start_question(self):
        '''Prepara a questao para ser tratada, zerando algumas variaveis de
        controle, colocando o cabecalho e aumentando o numero de questoes'''
        self.questionCounter = self.questionCounter + 1
        self.question_item = 'a'

        self.document.add_heading('Questao ' + str(self.questionCounter), level=2)
        p = self.document.add_paragraph('')

        self.paragraph = p

        # Flag de respostas False pois comecara com o enunciado
        self.respostas = False

# Parser Methods

    def handle_starttag(self, tag, attrs):
        # Adiciona imagem
        if tag == 'img':
            # Variaveis das imagens
            height = None
            width = None
            src = None
            # sai a procura das variaveis das imagens
            for attr in attrs:
                if attr[0] == 'src':
                    src = attr[1]
                elif attr[0] == 'style':
                    values = dict(item.split(":") for item in attr[1].split(";"))
                    for value in values:
                        if value.replace(' ', '') == 'width':
                            width = re.sub('\D', '', values[value])
                        elif value.replace(' ', '') == 'height':
                            height = re.sub('\D', '', values[value])
            # Coloca a imagem de acordo com sua respectiva dimensao definida (supondo DPI = 180)
            urllib.request.urlretrieve(src, "generateimage.png")
            if width:
                self.run.add_picture("generateimage.png", width=Inches(int(width)/180))
            else:
                self.run.add_picture("generateimage.png")
            print(src)
            os.remove("generateimage.png")
        # Habilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        elif tag == 'u':
            self.underline = True
        elif tag == 'em':
            self.italic = True
        elif tag == 'strong':
            self.bold = True

        # Textos de tabelas
        elif tag == 'p' and self.is_table:
            self.cell_paragraph = self.table.cell(
                        self.row_index, self.column_index).add_paragraph()
            self.cell_run = self.cell_paragraph.add_run()

        # Textos normais, nota que textos de respostas ocupam somente uma linha,
        # mesmo que contenham varias tags <p>
        elif tag == 'p' and not self.respostas:
            self.paragraph = self.document.add_paragraph()
            self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.run = self.paragraph.add_run()

        # Controle de tabelas
        elif tag == 'table':
            self.is_table = True
            self.table = self.document.add_table(rows=0, cols=0)
            self.table.style = 'TableGrid'
        elif tag == 'tr':
            self.table.add_row()
            self.row_index = self.row_index + 1
            self.column_index = -1
        elif tag == 'td':
            # Controle de adicao de tabelas, atualmente nao possui table merge
            if self.line_columns:
                self.table.add_column(1000)     # possivel bug do python-docx
            self.column_index = self.column_index + 1

    def handle_endtag(self, tag):
        # Desabilita variaveis que alteram os estilo do texto (sublinhado, negrito e italico)
        if tag == 'u':
            self.underline = False
        elif tag == 'em':
            self.italic = False
        elif tag == 'strong':
            self.bold = False

        # Desabilita os textos
        elif tag == 'p' and self.is_table:
            self.cell_run = None

        elif tag == 'p':
            self.run = None

        # Desabilita as tabelas
        elif tag == 'table':
            self.is_table = False
        elif tag == 'tr':
            self.line_columns = False
        # Desabilita as tabelas
        elif tag == 'td':
            self.cell_run = None
            # Exclui paragrafos vazios na celula atual, workaround encontrado
            # para melhorar o visual da lista gerada
            cell = self.table.cell(self.row_index, self.column_index)
            for para in cell.paragraphs:
                if para.text == "":
                    self.delete_paragraph(para)

    def handle_data(self, data):
        # Escreve o texto na tabela, verifica se o dado contem alguma coisa
        # importante, (existem linhas com tabulacoes que nao devem ser postas)
        if self.is_table and re.sub(r"\W", "", data) != "":
            # Texto normal sem tag
            if self.cell_run is None:
                self.cell_paragraph = self.table.cell(self.row_index,
                                self.column_index).add_paragraph()
                self.cell_paragraph.add_run(data)
            # Texto com tags
            else:
                self.cell_run.add_text(data)
                font = self.cell_run.font
                font.italic     = self.italic
                font.underline  = self.underline
                font.bold       = self.bold

        # Escreve o texto normalmente
        elif self.run is not None:
            self.run.add_text(data)
            font = self.run.font
            font.italic     = self.italic
            font.underline  = self.underline
            font.bold       = self.bold



class Answer_Parser():
    # Controle do documetno
    document = None
    docx_title = None

    def __init__(self, title):
        self.document = Document()
        self.docx_title = title

    def parse_heading(self, list_header):
        '''Cabecalho do gabarito gerado'''
        self.document.add_heading('Gabarito da lista' + list_header)

        self.document.add_paragraph(
            "Lista gerada em {:s} as {:s}.".format(
                datetime.date.today().strftime('%d/%m/%Y'),
                datetime.datetime.today().strftime('%X')
            )
        )

    def parse_answers_list_questions(self, list_questions):
        '''Faz o parser do gabarito, em formato de tabela'''

        question_counter = 0

        # Adiciona a tabela e seus cabecalhos
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Questao'
        hdr_cells[1].text = 'Reposta'

        # Faz o parser de cada questao
        for question in list_questions:
            question_counter = question_counter + 1
            table.add_row()
            table.cell(question_counter, 0).text = str(question_counter)

            question_item = 'a'
            answered = False
            # Procura pela questao correta
            for answer in question.answers.all():
                if answer.is_correct:
                    table.cell(question_counter, 1).text = question_item
                    answered = True
                question_item = chr(ord(question_item) + 1)

            # Se nao houve questao correta (NAO ERA PARA CAIR AQUI), fala sem
            # resposta
            if not answered:
                table.cell(question_counter, 1).text = 'Sem resposta'
